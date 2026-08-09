import streamlit as st
import streamlit.components.v1 as components
import ollama
import io
import json
import os
import sqlite3
import hashlib
import re
import html
import threading
import time
import uuid
from urllib.parse import quote, urlparse
from datetime import datetime
from pathlib import Path

# =========================================================
# CONFIG
# =========================================================
st.set_page_config(
    page_title="Nova",
    page_icon="✦",
    layout="wide",
    # Desktop opens the history rail; Streamlit automatically collapses it on
    # narrow phones so the conversation is visible first.
    initial_sidebar_state="auto",
)

TEMPERATURE = 0.5
WINDOW_MEMORY = 20


def get_secret(name):
    """Read a deployment secret without ever putting it in source code."""
    key = os.getenv(name, "").strip()
    if key:
        return key
    try:
        return str(st.secrets.get(name, "")).strip()
    except Exception:
        return ""


GEMINI_API_KEY = get_secret("GEMINI_API_KEY")
OPENROUTER_API_KEY = get_secret("OPENROUTER_API_KEY")
POLLINATIONS_API_KEY = get_secret("POLLINATIONS_API_KEY")
CLOUD_MODE = bool(GEMINI_API_KEY or OPENROUTER_API_KEY)

# A public Streamlit deployment must not share one SQLite file between users.
# Cloud chats are isolated per browser session; local Ollama keeps the existing
# persistent chat_history.db behaviour.
if CLOUD_MODE:
    if "nova_session_id" not in st.session_state:
        st.session_state.nova_session_id = uuid.uuid4().hex
    DB_FILE = os.path.join(
        "/tmp", f"nova_chat_{st.session_state.nova_session_id}.db"
    )
else:
    DB_FILE = "chat_history.db"

# =========================================================
# CHATGPT-LIKE CSS
# =========================================================
st.markdown("""
<style>
#MainMenu {visibility:hidden;}
footer {visibility:hidden;}
/* Keep the header shell alive because Streamlit places the reopen-sidebar
   control inside it. Hide only the toolbar/menu, not the whole header. */
header,
[data-testid="stHeader"] {
    visibility:visible !important;
    background:transparent !important;
}

[data-testid="stToolbar"],
[data-testid="stDecoration"] {
    display:none !important;
}

html, body, [class*="st-"], [data-testid="stAppViewContainer"] {
    font-family:-apple-system, BlinkMacSystemFont, "Segoe UI", Roboto,
        Helvetica, Arial, sans-serif !important;
}

:root, html, body {
    color-scheme:only light !important;
    background:#ffffff !important;
}

* {
    box-shadow:none !important;
}

[data-testid="stAppViewContainer"] {
    background:#ffffff !important;
    color:#202123 !important;
}

[data-testid="stMain"],
[data-testid="stBottom"],
[data-testid="stBottomBlockContainer"],
[data-testid="stBottom"] > div {
    background:#ffffff !important;
    color:#202123 !important;
}

/* Keep the existing conversation fully visible while Streamlit reruns. */
[data-stale="true"],
.stale {
    opacity:1 !important;
    filter:none !important;
    transition:none !important;
}

[data-testid="stAppViewContainer"],
[data-testid="stMain"],
[data-testid="stSidebar"] {
    filter:none !important;
}

[data-testid="stSidebar"] {
    background:#f7f7f8 !important;
    color:#202123 !important;
    border-right:1px solid #e5e5e5;
    overflow-x:hidden !important;
    width:310px !important;
    min-width:310px !important;
}

[data-testid="stSidebar"] > div:first-child {
    padding-top:1rem;
    overflow-x:hidden !important;
}

[data-testid="stSidebar"] p,
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] label,
[data-testid="stMarkdownContainer"] {
    color:#202123 !important;
}

[data-testid="stSidebar"] button,
[data-testid="stMain"] button {
    background:#ffffff !important;
    color:#202123 !important;
    border-color:#d9d9df !important;
}

textarea, input,
[data-baseweb="input"] > div,
[data-baseweb="textarea"] > div {
    background:#ffffff !important;
    color:#202123 !important;
    -webkit-text-fill-color:#202123 !important;
}

/* Reliable sidebar controls even when the Material Symbols font is blocked. */

[data-testid="collapsedControl"] button {
    display:flex !important;
    visibility:visible !important;
    align-items:center !important;
    justify-content:center !important;
    width:42px !important;
    height:42px !important;
    border:1px solid #dedede !important;
    border-radius:10px !important;
    background:#ffffff !important;
}

[data-testid="collapsedControl"] button span,
[data-testid="stSidebarCollapseButton"] button span {
    display:none !important;
}

[data-testid="collapsedControl"] button::after {
    content:"☰";
    font-family:Arial, sans-serif !important;
    font-size:24px !important;
    color:#303030 !important;
}

[data-testid="stSidebarCollapseButton"] button::after {
    content:"‹";
    font-family:Arial, sans-serif !important;
    font-size:28px !important;
    color:#555 !important;
}

/* Keep the history rail open on wide screens only. Phones use a drawer. */
@media (min-width:769px) {
    [data-testid="stSidebar"] {
        display:block !important;
        visibility:visible !important;
        transform:none !important;
        left:0 !important;
    }

    [data-testid="collapsedControl"],
    [data-testid="stSidebarCollapseButton"] {
        display:none !important;
    }
}

.block-container {
    max-width:1000px;
    padding-top:2rem;
    padding-bottom:7rem;
}

[data-testid="stChatMessage"] {
    border:none;
    padding-top:1rem;
    padding-bottom:1rem;
}

/* Streamlit sometimes exposes unloaded Material icon names such as
   "face" and "smart_toy" as text. Nova does not need chat avatars. */
[data-testid="stChatMessageAvatarUser"],
[data-testid="stChatMessageAvatarAssistant"],
[data-testid="stChatMessageAvatarCustom"],
[data-testid="stChatMessage"] [data-testid*="Avatar"] {
    display:none !important;
}

[data-testid="stChatMessageContent"] {
    width:100% !important;
    max-width:100% !important;
    color:#2f3037 !important;
    font-size:16px !important;
    line-height:1.65 !important;
}

/* ChatGPT-like Markdown rhythm for Nova answers. */
[data-testid="stChatMessageContent"] p {
    margin:0 0 0.9rem 0 !important;
}

[data-testid="stChatMessageContent"] h1,
[data-testid="stChatMessageContent"] h2,
[data-testid="stChatMessageContent"] h3 {
    color:#202123 !important;
    font-weight:650 !important;
    line-height:1.3 !important;
    margin:1.35rem 0 0.65rem !important;
}

[data-testid="stChatMessageContent"] h1 {font-size:1.55rem !important;}
[data-testid="stChatMessageContent"] h2 {font-size:1.3rem !important;}
[data-testid="stChatMessageContent"] h3 {font-size:1.12rem !important;}

[data-testid="stChatMessageContent"] ul,
[data-testid="stChatMessageContent"] ol {
    margin:0.35rem 0 1rem 1.4rem !important;
    padding-left:0.4rem !important;
}

[data-testid="stChatMessageContent"] li {
    margin:0.3rem 0 !important;
    padding-left:0.15rem !important;
}

[data-testid="stChatMessageContent"] pre {
    border:1px solid #e5e7eb !important;
    border-radius:10px !important;
    background:#f7f7f8 !important;
    padding:1rem !important;
    overflow-x:auto !important;
}

[data-testid="stChatMessageContent"] code:not(pre code) {
    background:#f1f1f3 !important;
    border-radius:5px !important;
    padding:0.12rem 0.32rem !important;
}

[data-testid="stChatMessageContent"] table {
    width:100% !important;
    border-collapse:collapse !important;
    margin:0.8rem 0 1.2rem !important;
}

[data-testid="stChatMessageContent"] th,
[data-testid="stChatMessageContent"] td {
    border:1px solid #dedede !important;
    padding:0.55rem 0.7rem !important;
    text-align:left !important;
}

[data-testid="stChatInput"] textarea {
    border-radius:24px !important;
    border:1px solid #d9d9e3 !important;
    background:#fff !important;
    padding-left:18px !important;
}

[data-testid="stChatInput"],
[data-testid="stChatInput"] > div,
[data-testid="stChatInput"] form {
    background:#f4f4f5 !important;
    color:#202123 !important;
    border-color:#e3e3e7 !important;
    forced-color-adjust:none !important;
}

[data-testid="stChatInput"] button {
    background:#e8eaed !important;
    color:#6b7280 !important;
    border:none !important;
}

.st-key-bottom_composer {
    width:min(1050px,92vw) !important;
    margin-left:auto !important;
    margin-right:auto !important;
    background:#ffffff !important;
}

.st-key-bottom_composer [data-testid="stHorizontalBlock"] {
    align-items:flex-end !important;
}

.st-key-bottom_composer [data-baseweb="select"] > div {
    border:1px solid #e4e4e7 !important;
    background:#f7f7f8 !important;
    color:#202123 !important;
    box-shadow:none !important;
    font-size:13px !important;
}

@media (max-width:900px) {
    .st-key-bottom_composer {
        width:96vw !important;
    }
}

/* Community Cloud branding can cover the composer inside mobile in-app
   browsers (LinkedIn, WhatsApp, etc.). */
[data-testid="stAppDeployButton"],
[class*="viewerBadge"],
[class*="ViewerBadge"] {
    display:none !important;
}

/* Mobile-first layout: the desktop history rail must never cover the chat.
   Visitors still get the current session and a full-width composer. */
@media (max-width:768px) {
    [data-testid="stSidebar"] {
        width:min(86vw, 310px) !important;
        min-width:min(86vw, 310px) !important;
    }

    [data-testid="collapsedControl"],
    [data-testid="stSidebarCollapseButton"] {
        display:flex !important;
        visibility:visible !important;
        z-index:1000001 !important;
    }

    [data-testid="stAppViewContainer"],
    [data-testid="stMain"],
    .main {
        margin-left:0 !important;
        width:100vw !important;
        max-width:100vw !important;
    }

    .block-container {
        width:100% !important;
        max-width:100% !important;
        padding:0.75rem 0.9rem 8rem !important;
    }

    .chat-title {
        margin-top:18vh !important;
        font-size:29px !important;
        line-height:1.25 !important;
        white-space:normal !important;
    }

    .st-key-bottom_composer {
        width:calc(100vw - 16px) !important;
        padding:0.4rem 0 !important;
        margin-bottom:3.4rem !important;
        background:#ffffff !important;
        forced-color-adjust:none !important;
    }

    .st-key-bottom_composer [data-testid="stHorizontalBlock"] {
        gap:0 !important;
    }

    .st-key-bottom_composer [data-testid="column"]:first-child {
        display:block !important;
        width:100% !important;
        flex:1 1 100% !important;
    }

    .st-key-bottom_composer [data-testid="column"]:last-child {
        display:none !important;
    }

    .st-key-bottom_composer [data-baseweb="select"],
    .st-key-bottom_composer [data-testid="stSelectbox"] {
        display:none !important;
    }

    [data-testid="stChatMessage"] {
        padding:0.7rem 0 !important;
    }

    [data-testid="stChatMessageContent"] {
        font-size:15px !important;
        line-height:1.55 !important;
    }

    [data-testid="stChatMessageContent"] table {
        display:block !important;
        max-width:100% !important;
        overflow-x:auto !important;
        white-space:normal !important;
    }
}

.chat-title {
    font-size:26px;
    font-weight:700;
    text-align:center;
    margin-top:90px;
}

.chat-subtitle {
    text-align:center;
    color:#777;
}

.file-chip {
    display:block;
    padding:7px 9px;
    margin:4px 0;
    background:#eeeeee;
    border-radius:8px;
    font-size:12px;
    overflow:hidden;
    white-space:nowrap;
    text-overflow:ellipsis;
}

.uploaded-file-note {
    display:inline-flex;
    align-items:center;
    gap:7px;
    margin-top:9px;
    padding:7px 11px;
    border:1px solid #dddddf;
    border-radius:10px;
    background:#ffffff;
    color:#444;
    font-size:13px;
}

.small-muted {
    color:#777;
    font-size:12px;
}
</style>
""", unsafe_allow_html=True)


# =========================================================
# SQLITE MEMORY
# =========================================================
def db():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS chats (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            chat_id INTEGER NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            created_at TEXT NOT NULL
        )
    """)
    conn.commit()
    return conn


def create_chat(title="New chat"):
    conn = db()
    now = datetime.now().isoformat(timespec="seconds")
    cur = conn.execute(
        "INSERT INTO chats(title, created_at, updated_at) VALUES(?,?,?)",
        (title, now, now),
    )
    chat_id = cur.lastrowid
    conn.commit()
    conn.close()
    return chat_id


def update_chat_title(chat_id, title):
    conn = db()
    conn.execute(
        "UPDATE chats SET title=?, updated_at=? WHERE id=?",
        (title[:60], datetime.now().isoformat(timespec="seconds"), chat_id),
    )
    conn.commit()
    conn.close()


def save_message(chat_id, role, content):
    conn = db()
    now = datetime.now().isoformat(timespec="seconds")
    conn.execute(
        "INSERT INTO messages(chat_id, role, content, created_at) VALUES(?,?,?,?)",
        (chat_id, role, content, now),
    )
    conn.execute(
        "UPDATE chats SET updated_at=? WHERE id=?",
        (now, chat_id),
    )
    conn.commit()
    conn.close()


def get_chats():
    conn = db()
    rows = conn.execute(
        "SELECT id,title FROM chats ORDER BY updated_at DESC"
    ).fetchall()
    conn.close()
    return rows


def get_messages(chat_id):
    conn = db()
    rows = conn.execute(
        "SELECT role,content FROM messages WHERE chat_id=? ORDER BY id",
        (chat_id,),
    ).fetchall()
    conn.close()
    return [{"role": r, "content": c} for r, c in rows]


def delete_chat(chat_id):
    conn = db()
    conn.execute("DELETE FROM messages WHERE chat_id=?", (chat_id,))
    conn.execute("DELETE FROM chats WHERE id=?", (chat_id,))
    conn.commit()
    conn.close()


# =========================================================
# SESSION STATE
# =========================================================
if "chat_id" not in st.session_state:
    chats = get_chats()
    st.session_state.chat_id = chats[0][0] if chats else create_chat()

if "messages" not in st.session_state:
    st.session_state.messages = get_messages(st.session_state.chat_id)

if "file_context" not in st.session_state:
    st.session_state.file_context = ""

if "uploaded_files" not in st.session_state:
    st.session_state.uploaded_files = []

if "file_signature" not in st.session_state:
    st.session_state.file_signature = ""

if "generated_pdf" not in st.session_state:
    st.session_state.generated_pdf = None

if "generated_pdf_name" not in st.session_state:
    st.session_state.generated_pdf_name = "Nova_Chat.pdf"

if "generated_pdf_chat_id" not in st.session_state:
    st.session_state.generated_pdf_chat_id = None


# =========================================================
# PDF EXPORT
# =========================================================
def get_pdf_request_mode(query, has_uploaded_file=False):
    """Return 'content', 'export', or None for PDF-related requests."""
    q = query.lower()

    # A follow-up asking for the missing link should export the conversation.
    if re.search(r"\b(download\s*(link|button)|where.*download)\b", q):
        return "export"

    if "pdf" not in q:
        return None

    # When a source file is already attached, a follow-up such as "PDF bana
    # do" means create the requested document from that file. Only an explicit
    # request for the chat/conversation itself should export chat history.
    explicit_chat_export = [
        "chat pdf", "chat ka pdf", "conversation pdf",
        "export chat", "previous messages", "chat history",
    ]
    if has_uploaded_file and not any(x in q for x in explicit_chat_export):
        return "content"

    content_terms = [
        "question", "questions", "answer", "answers", "q&a", "qna",
        "interview", "job description", " jd ", "resume", "cv", "report",
        "modify", "update", "rewrite", "ats",
        "guide", "notes", "summary", "relevant", "sare", "saare",
    ]
    request_terms = [
        "make", "create", "generate", "prepare", "want", "need",
        "chahiye", "bana", "banao", "banado", "de do", "dede",
    ]

    if any(term in f" {q} " for term in content_terms) and any(
        term in q for term in request_terms
    ):
        return "content"

    export_terms = [
        "make", "create", "generate", "prepare", "export", "convert",
        "download", "bana", "banao", "banado", "chahiye",
    ]
    if any(term in q for term in export_terms):
        return "export"

    return None


def wants_pdf_generation(query):
    return get_pdf_request_mode(query) is not None


def is_resume_request(query):
    """Detect CV work across the current request and recent conversation."""
    recent_user_text = " ".join(
        str(message.get("content", ""))
        for message in st.session_state.get("messages", [])[-8:]
        if message.get("role") == "user"
    )
    filenames = " ".join(st.session_state.get("uploaded_files", []))
    combined = f"{recent_user_text} {query} {filenames}".lower()
    return any(term in combined for term in [
        " resume", "resume ", " cv", "cv ", "curriculum vitae"
    ])


def build_chat_pdf(messages, title="Nova - Chat Export", show_role_labels=True):
    # ReportLab is imported only when a PDF is requested, keeping startup fast.
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError as exc:
        raise RuntimeError(
            "PDF support is not installed. Run: pip install reportlab"
        ) from exc

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title,
        author="Nova",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ChatTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#202123"),
        alignment=TA_CENTER,
        spaceAfter=8 * mm,
    )
    user_style = ParagraphStyle(
        "UserLabel",
        parent=styles["Heading3"],
        textColor=colors.HexColor("#C2414D"),
        spaceBefore=4 * mm,
        spaceAfter=2 * mm,
    )
    assistant_style = ParagraphStyle(
        "AssistantLabel",
        parent=styles["Heading3"],
        textColor=colors.HexColor("#C88912"),
        spaceBefore=4 * mm,
        spaceAfter=2 * mm,
    )
    body_style = ParagraphStyle(
        "ChatBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor("#303030"),
        spaceAfter=2.5 * mm,
    )

    story = [
        Paragraph(html.escape(title), title_style),
        Paragraph(
            "Generated on " + datetime.now().strftime("%d %B %Y, %I:%M %p"),
            styles["Normal"],
        ),
        Spacer(1, 5 * mm),
    ]

    for message in messages:
        role = message.get("role", "assistant")
        label = "You" if role == "user" else "Nova"
        label_style = user_style if role == "user" else assistant_style
        if show_role_labels:
            story.append(Paragraph(label, label_style))

        content = str(message.get("content", "")).strip()
        blocks = re.split(r"\n\s*\n", content) if content else [""]
        for block in blocks:
            safe = html.escape(block).replace("\n", "<br/>")
            safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
            story.append(Paragraph(safe or " ", body_style))

        story.append(Spacer(1, 2 * mm))

    def add_page_number(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawCentredString(
            A4[0] / 2,
            9 * mm,
            f"Page {document.page}",
        )
        canvas.restoreState()

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    output.seek(0)
    return output.getvalue()


def build_resume_pdf(content):
    """Create a clean multi-page ATS resume instead of a chat export."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    )

    output = io.BytesIO()
    doc = SimpleDocTemplate(
        output, pagesize=A4, rightMargin=15 * mm, leftMargin=15 * mm,
        topMargin=13 * mm, bottomMargin=13 * mm,
        title="Tanya Tiwari Resume", author="Tanya Tiwari",
    )
    base = getSampleStyleSheet()
    name_style = ParagraphStyle(
        "ResumeName", parent=base["Title"], fontName="Helvetica-Bold",
        fontSize=20, leading=23, alignment=TA_CENTER,
        textColor=colors.HexColor("#202124"), spaceAfter=3 * mm,
    )
    section_style = ParagraphStyle(
        "ResumeSection", parent=base["Heading2"], fontName="Helvetica-Bold",
        fontSize=11.5, leading=14, textColor=colors.HexColor("#202124"),
        spaceBefore=4 * mm, spaceAfter=1.2 * mm,
        keepWithNext=True,
    )
    role_style = ParagraphStyle(
        "ResumeRole", parent=base["Heading3"], fontName="Helvetica-Bold",
        fontSize=10.2, leading=13, textColor=colors.HexColor("#25262b"),
        spaceBefore=2.2 * mm, spaceAfter=0.8 * mm, keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "ResumeBody", parent=base["BodyText"], fontName="Helvetica",
        fontSize=9.5, leading=13, textColor=colors.HexColor("#303030"),
        spaceAfter=1.2 * mm,
    )
    contact_style = ParagraphStyle(
        "ResumeContact", parent=body_style, alignment=TA_CENTER,
        fontSize=9.2, leading=12, spaceAfter=1 * mm,
    )
    bullet_style = ParagraphStyle(
        "ResumeBullet", parent=body_style, leftIndent=4 * mm,
        firstLineIndent=-3 * mm, spaceAfter=1 * mm,
    )

    # Normalize common mojibake and typography before ReportLab rendering.
    content = str(content)
    for broken in ["â■■", "â–■■", "â€“", "â€”", "–", "—"]:
        content = content.replace(broken, "-")
    story = []
    first_content = True
    before_first_section = True
    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            story.append(Spacer(1, 1.1 * mm))
            continue

        heading_match = re.match(r"^(#{1,3})\s*(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading = heading_match.group(2).replace("**", "")
            safe_heading = html.escape(heading)
            if first_content or level == 1:
                story.append(Paragraph(safe_heading.upper(), name_style))
                first_content = False
            elif level == 2:
                before_first_section = False
                story.append(Spacer(1, 1.2 * mm))
                story.append(Paragraph(safe_heading.upper(), section_style))
                story.append(HRFlowable(
                    width="100%", thickness=0.55,
                    color=colors.HexColor("#5f6368"),
                    spaceBefore=0, spaceAfter=1.5 * mm,
                ))
            else:
                before_first_section = False
                story.append(Paragraph(safe_heading, role_style))
            continue

        is_bullet = bool(re.match(r"^[-*•]\s+", line))
        clean = re.sub(r"^[-*•]\s+", "", line)
        clean = re.sub(r"(?<!\*)\*(?!\*)", "", clean)
        safe = html.escape(clean)
        safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
        if first_content:
            story.append(Paragraph(safe.upper(), name_style))
            first_content = False
        elif is_bullet:
            story.append(Paragraph("• " + safe, bullet_style))
        elif before_first_section:
            story.append(Paragraph(safe, contact_style))
        else:
            story.append(Paragraph(safe, body_style))

    def page_number(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawCentredString(A4[0] / 2, 7 * mm, str(document.page))
        canvas.restoreState()

    doc.build(story, onFirstPage=page_number, onLaterPages=page_number)
    output.seek(0)
    return output.getvalue()


# =========================================================
# FILE READERS
# =========================================================
def read_pdf(file):
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file.getvalue()))
        return "\n".join(
            f"--- PDF Page {i+1} ---\n{page.extract_text() or ''}"
            for i, page in enumerate(reader.pages)
        )
    except Exception as e:
        return f"Could not read PDF: {e}"


def read_excel(file):
    try:
        import pandas as pd

        raw = file.getvalue()
        book = pd.ExcelFile(io.BytesIO(raw))
        parts = []
        for sheet in book.sheet_names:
            df = pd.read_excel(io.BytesIO(raw), sheet_name=sheet)
            parts.append(f"--- Excel Sheet: {sheet} ---")
            parts.append(df.to_string(index=False))
        return "\n".join(parts)
    except Exception as e:
        return f"Could not read Excel: {e}"


def read_csv(file):
    try:
        import pandas as pd

        raw = file.getvalue()
        for enc in ["utf-8", "utf-8-sig", "latin-1"]:
            try:
                return pd.read_csv(
                    io.BytesIO(raw), encoding=enc
                ).to_string(index=False)
            except Exception:
                pass
        return "Could not decode CSV."
    except Exception as e:
        return f"Could not read CSV: {e}"


def read_docx(file):
    try:
        from docx import Document

        doc = Document(io.BytesIO(file.getvalue()))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception as e:
        return f"Could not read DOCX: {e}"


def read_text(file):
    raw = file.getvalue()
    for enc in ["utf-8", "utf-8-sig", "latin-1"]:
        try:
            return raw.decode(enc)
        except Exception:
            pass
    return "Could not decode text."


def read_json(file):
    try:
        return json.dumps(
            json.loads(file.getvalue().decode("utf-8")),
            indent=2,
            ensure_ascii=False,
        )
    except Exception as e:
        return f"Could not read JSON: {e}"


def read_binary(file):
    try:
        return file.getvalue().decode("utf-8")
    except Exception:
        return (
            "Binary file uploaded. Raw binary content cannot be directly "
            "read as text by this version."
        )


def extract_file(file):
    ext = os.path.splitext(file.name.lower())[1]

    if ext == ".pdf":
        return read_pdf(file)
    if ext in [".xlsx", ".xls", ".xlsm"]:
        return read_excel(file)
    if ext == ".csv":
        return read_csv(file)
    if ext == ".docx":
        return read_docx(file)
    if ext == ".json":
        return read_json(file)
    if ext in [
        ".txt", ".md", ".py", ".sql", ".sas", ".r",
        ".html", ".css", ".js", ".xml", ".yaml",
        ".yml", ".log"
    ]:
        return read_text(file)

    return read_binary(file)


def make_signature(files):
    return "|".join(
        f"{f.name}:{len(f.getvalue())}:"
        f"{hashlib.md5(f.getvalue()).hexdigest()}"
        for f in files
    )


def build_file_context(files):
    parts = []

    for f in files:
        content = extract_file(f)

        # Avoid sending enormous files directly into the model context.
        if len(content) > 60000:
            content = (
                content[:60000]
                + "\n[File content truncated for model context.]"
            )

        parts.append(
            f"================ FILE: {f.name} ================\n"
            f"{content}"
        )

    return "\n\n".join(parts)


def should_use_file_context(query):
    """Attach uploaded content only when the question refers to that content."""
    q = query.lower()
    file_terms = [
        "file", "uploaded", "upload", "attached", "attachment",
        "pdf", "excel", "csv", "dataset", "document", "docx",
        "job description", " jd ", "this data", "this sheet",
        "this document", "isme", "iske", "iss file", "is file",
    ]
    padded = f" {q} "
    return any(term in padded for term in file_terms)


# =========================================================
# FAST WEB ROUTER + DOWNLOAD TOOL
# =========================================================
def needs_web_search(query):
    q = query.lower()

    # Any year after the model's knowledge cutoff should automatically use
    # live search (for example: 2024, 2025, 2026...).
    years = [int(y) for y in re.findall(r"\b(20\d{2})\b", q)]
    if any(year >= 2024 for year in years):
        return True

    keywords = [
        "latest", "today", "current", "now", "news", "recent",
        "this year", "this month", "this week", "yesterday", "tomorrow",
        "price", "weather", "stock", "share price", "score", "result",
        "schedule", "election", "president", "prime minister", "ceo",
        "salary", "package", "hike", "vacancy", "job opening", "release",
        "career", "careers", "jobs page", "career page", "job listing",
        "version", "law", "rule", "policy", "rate", "exchange rate",
        "market", "search internet", "search web", "on internet",
        "google", "online", "website", "download", "csv", "xlsx",
        "excel", "dataset", "data file",
        # Exact factual/statistical questions must be verified instead of
        # answered from model memory.
        "census", "population", "according to", "statistic", "statistics",
        "percentage", "how many", "total number", "age group", "literacy",
        "gdp", "inflation", "unemployment", "official data", "report says",
    ]
    return any(x in q for x in keywords)


def wants_weather(query):
    q = query.lower()
    return any(term in q for term in [
        "weather", "temperature", "forecast", "mausam", "tapman",
        "बारिश", "मौसम", "तापमान",
    ])


def extract_weather_location(query):
    """Extract a likely city/place from English or Hinglish weather prompts."""
    cleaned = str(query)
    cleaned = re.sub(r"([A-Za-z]+)['’]s\b", r"\1", cleaned)
    patterns = [
        r"\b(?:please|pls|current|live|today(?:'s)?|now|update|tell me|show)\b",
        r"\b(?:weather|temperature|forecast|mausam|tapman|degree|degrees)\b",
        r"\b(?:in|for|of|ka|ki|k|the|is|what|whats|please)\b",
    ]
    for pattern in patterns:
        cleaned = re.sub(pattern, " ", cleaned, flags=re.I)
    cleaned = re.sub(r"[^\w\s\-]", " ", cleaned, flags=re.UNICODE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned or str(query).strip()


def get_live_weather(query):
    """Return current conditions without consuming any LLM quota."""
    import requests

    location = extract_weather_location(query)
    geo_response = requests.get(
        "https://geocoding-api.open-meteo.com/v1/search",
        params={"name": location, "count": 1, "language": "en", "format": "json"},
        timeout=12,
    )
    geo_response.raise_for_status()
    places = geo_response.json().get("results") or []
    if not places:
        raise RuntimeError(f"I couldn't find the location '{location}'.")

    place = places[0]
    weather_response = requests.get(
        "https://api.open-meteo.com/v1/forecast",
        params={
            "latitude": place["latitude"],
            "longitude": place["longitude"],
            "current": (
                "temperature_2m,apparent_temperature,relative_humidity_2m,"
                "precipitation,weather_code,wind_speed_10m"
            ),
            "daily": "temperature_2m_max,temperature_2m_min,precipitation_probability_max",
            "timezone": "auto",
            "forecast_days": 1,
        },
        timeout=12,
    )
    weather_response.raise_for_status()
    data = weather_response.json()
    current = data.get("current") or {}
    daily = data.get("daily") or {}

    weather_codes = {
        0: "Clear sky", 1: "Mainly clear", 2: "Partly cloudy", 3: "Overcast",
        45: "Fog", 48: "Rime fog", 51: "Light drizzle", 53: "Drizzle",
        55: "Heavy drizzle", 61: "Light rain", 63: "Rain", 65: "Heavy rain",
        71: "Light snow", 73: "Snow", 75: "Heavy snow", 80: "Rain showers",
        81: "Rain showers", 82: "Heavy rain showers", 95: "Thunderstorm",
        96: "Thunderstorm with hail", 99: "Heavy thunderstorm with hail",
    }
    code = int(current.get("weather_code", -1))
    condition = weather_codes.get(code, "Current conditions")
    place_name = place.get("name", location)
    region = place.get("admin1") or place.get("country") or ""
    full_place = f"{place_name}, {region}" if region else place_name

    def first(values, default="N/A"):
        return values[0] if isinstance(values, list) and values else default

    return (
        f"### Weather in {full_place}\n\n"
        f"- **Current temperature:** {current.get('temperature_2m', 'N/A')}°C\n"
        f"- **Feels like:** {current.get('apparent_temperature', 'N/A')}°C\n"
        f"- **Condition:** {condition}\n"
        f"- **Humidity:** {current.get('relative_humidity_2m', 'N/A')}%\n"
        f"- **Wind:** {current.get('wind_speed_10m', 'N/A')} km/h\n"
        f"- **Today's range:** {first(daily.get('temperature_2m_min'))}°C to "
        f"{first(daily.get('temperature_2m_max'))}°C\n"
        f"- **Maximum rain chance:** "
        f"{first(daily.get('precipitation_probability_max'))}%\n\n"
        f"*Updated {current.get('time', 'now')} local time. Weather data by "
        f"[Open-Meteo](https://open-meteo.com/).*"
    )


def needs_fact_verification(query):
    """Identify numeric/public-fact questions that need authoritative sources."""
    q = query.lower()
    terms = [
        "census", "population", "according to", "statistics", "percentage",
        "how many", "total number", "age group", "literacy", "gdp",
        "inflation", "unemployment", "official", "report",
    ]
    return any(term in q for term in terms) or bool(
        re.search(r"\b(?:19|20)\d{2}\b", q)
    )


def is_code_build_request(query):
    """Route app/website/clone requests to a complete runnable-code prompt."""
    q = query.lower()
    build_terms = [
        "build", "create", "make", "develop", "generate", "banao", "bana do",
        "bnado", "clone", "website", "web app", "application", "dashboard",
        "frontend", "backend", "full stack", "streamlit app", "project code",
    ]
    code_terms = [
        "code", "python", "html", "css", "javascript", "react", "streamlit",
        "sql", "api", "spotify", "netflix", "portfolio", "todo app",
    ]
    return any(term in q for term in build_terms) and any(
        term in q for term in code_terms
    )


def wants_download(query):
    q = query.lower()
    return any(x in q for x in [
        "download", "save the file", "get the file", "fetch the file",
        "csv", "xlsx", "excel file", "dataset"
    ])


def extract_urls(text):
    return re.findall(r'https?://[^\s<>"\']+', text or "")


def web_search(query):
    try:
        try:
            from ddgs import DDGS
        except ImportError:
            from duckduckgo_search import DDGS

        results = DDGS().text(query, max_results=5)
        return "\n\n".join(
            f"Title: {r.get('title','')}\n"
            f"URL: {r.get('href','')}\n"
            f"Summary: {r.get('body','')}"
            for r in results
        ) or "No web results found."
    except Exception as e:
        return f"Web search unavailable: {e}"


def build_search_fallback(query, raw_results):
    """Present useful live sources when every LLM provider is temporarily busy."""
    if not raw_results or raw_results.startswith("Web search unavailable:"):
        return ""

    items = []
    for block in raw_results.split("\n\n"):
        title = re.search(r"^Title:\s*(.+)$", block, flags=re.M)
        url = re.search(r"^URL:\s*(.+)$", block, flags=re.M)
        summary = re.search(r"^Summary:\s*(.*)$", block, flags=re.M | re.S)
        if not title or not url:
            continue
        label = title.group(1).strip()
        href = url.group(1).strip()
        text = summary.group(1).strip() if summary else ""
        items.append(f"- **[{label}]({href})**\n  {text}")

    if not items:
        return ""
    return (
        "The AI model is temporarily busy, but I found these live results for "
        f"**{query.strip()}**:\n\n" + "\n\n".join(items[:5])
        + "\n\n*Open a source above to verify the exact detail.*"
    )


def wants_images(query):
    """Recognize direct requests for pictures in English or Hinglish."""
    q = f" {query.lower()} "
    image_terms = [
        " image ", " images ", " pic ", " pics ", " picture ",
        " pictures ", " photo ", " photos ", " wallpaper ",
        " poster ", " logo ", " artwork ", " art ",
        " tasveer ", " tasveeren ", " photo dikhao ", " pics dikhao ",
    ]
    return any(term in q for term in image_terms)


def wants_image_generation(query):
    """Separate AI creation from ordinary web image search."""
    q = query.lower()
    image_word = any(term in q for term in [
        "image", "picture", "photo", "pic", "wallpaper", "poster",
        "logo", "art", "tasveer",
    ])
    creation_word = any(term in q for term in [
        "generate", "create", "make", "draw", "design", "banao",
        "bana do", "bnao", "bnado", "generate karo", "create karo",
    ])
    return image_word and creation_word


def wants_image_search(query):
    """Use web photos only when the user clearly asks for real/search results."""
    q = query.lower()
    return any(term in q for term in [
        "search", "find", "from web", "from google", "google se",
        "real photo", "real image", "actual photo", "latest photo",
        "current photo", "news photo", "photos dikhao", "pics dikhao",
        "photo dikhao", "images dikhao",
    ])


def requests_broad_scope(query):
    """Detect when the latest question intentionally drops an old filter."""
    q = query.lower()
    return any(term in q for term in [
        "all positions", "all jobs", "all companies", "every company",
        "any company", "overall", "in general", "generally",
        "jitni bhi", "sabhi position", "sab position", "sb position",
        "sab dikhao", "sb dikhao", "saari position", "sari position",
    ])


def generate_ai_image(query):
    """Generate and cache an image using Pollinations' image endpoint."""
    import requests

    prompt = re.sub(
        r"\b(?:please|generate|create|make|draw|design|banao|bana\s+do|"
        r"bnao|bnado|karo|an?|the|image|picture|photo|pic)\b",
        " ",
        query,
        flags=re.I,
    )
    prompt = re.sub(r"\s+", " ", prompt).strip() or query.strip()
    seed = int(hashlib.sha256(query.encode("utf-8")).hexdigest()[:8], 16)
    url = "https://gen.pollinations.ai/image/" + quote(prompt, safe="")
    params = {
        "width": 1024,
        "height": 1024,
        "seed": seed,
        "nologo": "true",
    }
    if POLLINATIONS_API_KEY:
        params["key"] = POLLINATIONS_API_KEY

    response = requests.get(url, params=params, timeout=(20, 240))
    response.raise_for_status()
    content_type = response.headers.get("Content-Type", "").lower()
    if "image" not in content_type or len(response.content) < 5000:
        raise RuntimeError("The image service did not return a valid image.")

    session_id = st.session_state.get("nova_session_id", "local")
    image_path = os.path.join(
        "/tmp", f"nova_generated_{session_id}_{uuid.uuid4().hex}.jpg"
    )
    Path(image_path).write_bytes(response.content)
    return prompt, image_path


def image_search(query, max_results=6):
    """Return real image results instead of asking the LLM to imagine them."""
    try:
        from ddgs import DDGS
    except ImportError:
        from duckduckgo_search import DDGS

    cleaned_query = re.sub(
        r"\b(images?|pics?|pictures?|photos?|wallpapers?|dikhao|show me)\b",
        " ",
        query,
        flags=re.I,
    )
    cleaned_query = re.sub(r"\s+", " ", cleaned_query).strip() or query

    results = DDGS().images(
        cleaned_query,
        region="wt-wt",
        safesearch="moderate",
        max_results=max_results,
    )
    images = []
    for item in results or []:
        image_url = item.get("image")
        if not image_url:
            continue
        images.append({
            "image": image_url,
            "thumbnail": item.get("thumbnail") or image_url,
            "title": item.get("title") or cleaned_query.title(),
            "source": item.get("url") or "",
        })
    return images


IMAGE_RESULT_PREFIX = "[[NOVA_IMAGE_RESULTS]]"
GENERATED_IMAGE_PREFIX = "[[NOVA_GENERATED_IMAGE]]"


def pack_image_results(query, images, clean_gallery=False):
    return IMAGE_RESULT_PREFIX + json.dumps(
        {
            "query": query,
            "images": images,
            "clean_gallery": clean_gallery,
        },
        ensure_ascii=False,
    )


def pack_generated_image(prompt, image_path):
    return GENERATED_IMAGE_PREFIX + json.dumps(
        {"prompt": prompt, "path": image_path}, ensure_ascii=False
    )


def repair_mojibake(value):
    """Repair UTF-8 text that an SSE client mistakenly decoded as Latin-1."""
    text = str(value or "")
    suspicious = ("Ã", "Â", "â€", "à¤", "à¥", "ðŸ")
    if not any(marker in text for marker in suspicious):
        return text
    try:
        repaired = text.encode("latin-1").decode("utf-8")
        return repaired if repaired else text
    except (UnicodeEncodeError, UnicodeDecodeError):
        return text


def render_chat_content(content):
    """Render normal messages or a persistent grid of image-search results."""
    content = repair_mojibake(content)
    if content.startswith(GENERATED_IMAGE_PREFIX):
        try:
            payload = json.loads(content[len(GENERATED_IMAGE_PREFIX):])
            image_path = payload["path"]
            st.markdown(
                f"Generated image for **{payload.get('prompt', 'your prompt')}**:"
            )
            if os.path.exists(image_path):
                st.image(image_path, use_container_width=True)
                st.download_button(
                    "⬇️ Download image",
                    data=Path(image_path).read_bytes(),
                    file_name="Nova_Generated_Image.jpg",
                    mime="image/jpeg",
                    key="generated_" + hashlib.md5(
                        image_path.encode("utf-8")
                    ).hexdigest(),
                )
            else:
                st.info("This temporary generated image has expired. Generate it again.")
        except Exception:
            st.warning("The generated image could not be displayed.")
        return

    if not content.startswith(IMAGE_RESULT_PREFIX):
        if "\n\n📎 " in content:
            body, attached = content.rsplit("\n\n📎 ", 1)
            st.markdown(body)
            st.markdown(
                '<div class="uploaded-file-note">📎 '
                + html.escape(attached)
                + '</div>',
                unsafe_allow_html=True,
            )
        else:
            st.markdown(content)
        return

    try:
        payload = json.loads(content[len(IMAGE_RESULT_PREFIX):])
        images = payload.get("images", [])
        clean_gallery = bool(payload.get("clean_gallery"))
        st.markdown(f"Here are pictures for **{payload.get('query', 'your search')}**:")
        for start in range(0, len(images), 3):
            columns = st.columns(3)
            for column, item in zip(columns, images[start:start + 3]):
                with column:
                    st.image(
                        item.get("thumbnail") or item["image"],
                        use_container_width=True,
                    )
                    if not clean_gallery:
                        title = item.get("title", "View image")
                        source = item.get("source", "")
                        if source:
                            st.markdown(f"[{title}]({source})")
                        else:
                            st.caption(title)
    except Exception:
        st.error("These image results could not be displayed.")


def download_file(url, folder="downloads"):
    import requests

    os.makedirs(folder, exist_ok=True)
    url = url.rstrip(".,);]}>")
    name = os.path.basename(urlparse(url).path) or "downloaded_file"
    name = re.sub(r"[^A-Za-z0-9._-]", "_", name)

    try:
        response = requests.get(
            url,
            timeout=25,
            allow_redirects=True,
            headers={"User-Agent": "Nova/1.0"}
        )
        response.raise_for_status()

        content_type = response.headers.get("content-type", "").lower()

        if "." not in name:
            if "csv" in content_type:
                name += ".csv"
            elif "spreadsheet" in content_type or "excel" in content_type:
                name += ".xlsx"
            elif "json" in content_type:
                name += ".json"

        path = os.path.join(folder, name)

        with open(path, "wb") as f:
            f.write(response.content)

        return path, f"Downloaded successfully: {path}"

    except Exception as e:
        return None, f"Download failed: {e}"


def try_download_from_results(search_text):
    urls = extract_urls(search_text)

    preferred = [
        u for u in urls
        if re.search(r'\.(csv|xlsx|xls|json)(\?|$)', u, re.I)
    ]

    candidates = preferred + [u for u in urls if u not in preferred]

    for url in candidates[:8]:
        path, status = download_file(url)
        if path:
            return path, status

    return None, "No directly downloadable public data file was found."


# =========================================================
# HYBRID AI BACKEND: GEMINI CLOUD + OPTIONAL OPENROUTER + LOCAL OLLAMA
# =========================================================
OPENROUTER_MODEL_CANDIDATES = {
    "llama3.2": [
        "meta-llama/llama-3.3-70b-instruct",
        "meta-llama/llama-3.1-8b-instruct",
        "openrouter/auto",
    ],
    "deepseek-r1:latest": [
        "~deepseek/deepseek-v4-flash-latest",
        "deepseek/deepseek-v3.2",
        "deepseek/deepseek-r1",
        "openrouter/auto",
    ],
    "qwen2.5:latest": [
        "qwen/qwen3.7-flash",
        "qwen/qwen3-30b-a3b-instruct-2507",
        "qwen/qwen-2.5-72b-instruct",
        "openrouter/auto",
    ],
}


class ApiRequestGate:
    """Smooth bursts from many Streamlit sessions sharing one free API key."""

    def __init__(self, minimum_interval=1.1):
        self.minimum_interval = minimum_interval
        self.lock = threading.Lock()
        self.last_request = 0.0

    def wait_turn(self):
        with self.lock:
            remaining = self.minimum_interval - (time.monotonic() - self.last_request)
            if remaining > 0:
                time.sleep(remaining)
            self.last_request = time.monotonic()


@st.cache_resource(show_spinner=False)
def get_api_request_gate():
    return ApiRequestGate()


@st.cache_data(ttl=3600, show_spinner=False)
def get_openrouter_model_ids():
    """Cache model availability so changed/deprecated IDs do not break Nova."""
    import requests

    response = requests.get(
        "https://openrouter.ai/api/v1/models",
        timeout=12,
    )
    response.raise_for_status()
    return {item["id"] for item in response.json().get("data", [])}


def resolve_openrouter_model(local_model):
    candidates = OPENROUTER_MODEL_CANDIDATES.get(
        local_model, ["openrouter/auto"]
    )
    try:
        available = get_openrouter_model_ids()
        return next((name for name in candidates if name in available), candidates[0])
    except Exception:
        return candidates[0]


def stream_openrouter_chat(local_model, messages, max_tokens):
    """Yield OpenRouter text deltas in the same shape used by the UI."""
    import requests

    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json",
            "X-Title": "Nova",
        },
        json={
            "model": resolve_openrouter_model(local_model),
            "messages": messages,
            "temperature": TEMPERATURE,
            "max_tokens": max_tokens,
            "stream": True,
        },
        timeout=(15, 180),
        stream=True,
    )
    response.raise_for_status()

    # SSE responses often omit a charset. requests then assumes Latin-1 and
    # corrupts Hindi/emoji. Always decode the complete event line as UTF-8.
    for raw_line in response.iter_lines(decode_unicode=False):
        if isinstance(raw_line, bytes):
            raw_line = raw_line.decode("utf-8", errors="replace")
        if not raw_line or not raw_line.startswith("data:"):
            continue
        payload = raw_line[5:].strip()
        if payload == "[DONE]":
            break
        try:
            data = json.loads(payload)
            piece = data.get("choices", [{}])[0].get("delta", {}).get(
                "content", ""
            )
            if piece:
                yield {"message": {"content": piece}}
        except (json.JSONDecodeError, IndexError, TypeError):
            continue


def stream_gemini_chat(
    model_name, messages, max_tokens, temperature=TEMPERATURE
):
    """Stream Gemini output using REST, so no extra Google SDK is required."""
    import requests

    system_text = "\n\n".join(
        message.get("content", "")
        for message in messages
        if message.get("role") == "system"
    )
    contents = []
    for message in messages:
        role = message.get("role")
        if role == "system":
            continue
        contents.append({
            "role": "model" if role == "assistant" else "user",
            "parts": [{"text": message.get("content", "")}],
        })

    # Google periodically changes which models receive free-tier quota. Try
    # current models first and retain 2.5 as compatibility fallbacks.
    model_candidates = list(dict.fromkeys([
        model_name,
        "gemini-3.6-flash",
        "gemini-3-flash-preview",
        "gemini-3.1-flash-lite",
        "gemini-2.5-flash",
        "gemini-2.5-flash-lite",
    ]))
    errors = []
    for candidate in model_candidates:
        response = None
        for attempt in range(2):
            get_api_request_gate().wait_turn()
            response = requests.post(
                "https://generativelanguage.googleapis.com/v1beta/models/"
                f"{candidate}:streamGenerateContent?alt=sse",
                headers={
                    "x-goog-api-key": GEMINI_API_KEY,
                    "Content-Type": "application/json",
                },
                json={
                    "systemInstruction": {"parts": [{"text": system_text}]},
                    "contents": contents,
                    "generationConfig": {
                        "temperature": temperature,
                        "maxOutputTokens": max_tokens,
                    },
                },
                timeout=(15, 180),
                stream=True,
            )
            if response.ok:
                break

            status = response.status_code
            errors.append(f"{candidate}: HTTP {status}")
            retryable = status in (429, 500, 503, 504)
            if retryable and attempt == 0:
                try:
                    wait_seconds = float(response.headers.get("Retry-After", "1.5"))
                except ValueError:
                    wait_seconds = 1.5
                time.sleep(min(max(wait_seconds, 0.8), 4.0))
                continue
            break

        if response is None or not response.ok:
            # Invalid/authenticated requests will not improve by trying more
            # models with the same key.
            if response is not None and response.status_code in (401, 403):
                break
            continue

        produced_text = False
        # Gemini SSE may not declare UTF-8 in Content-Type. Explicit decoding
        # is required for Hindi and every other non-ASCII script.
        for raw_line in response.iter_lines(decode_unicode=False):
            if isinstance(raw_line, bytes):
                raw_line = raw_line.decode("utf-8", errors="replace")
            if not raw_line or not raw_line.startswith("data:"):
                continue
            try:
                data = json.loads(raw_line[5:].strip())
                parts = (
                    data.get("candidates", [{}])[0]
                    .get("content", {})
                    .get("parts", [])
                )
                piece = "".join(
                    part.get("text", "") for part in parts if "text" in part
                )
                if piece:
                    produced_text = True
                    yield {"message": {"content": piece}}
            except (json.JSONDecodeError, IndexError, TypeError):
                continue

        if produced_text:
            return
        errors.append(f"{candidate}: empty response")

    raise RuntimeError("Gemini unavailable: " + "; ".join(errors[-8:]))


def stream_nova_chat(local_model, messages, model_options):
    errors = []
    max_tokens = model_options.get("num_predict", 1100)

    if GEMINI_API_KEY:
        try:
            produced = False
            for chunk in stream_gemini_chat(
                local_model,
                messages,
                max_tokens,
                model_options.get("temperature", TEMPERATURE),
            ):
                produced = True
                yield chunk
            if produced:
                return
        except Exception as exc:
            errors.append(f"Gemini: {exc}")

    if OPENROUTER_API_KEY:
        try:
            produced = False
            for chunk in stream_openrouter_chat(
                local_model, messages, max_tokens
            ):
                produced = True
                yield chunk
            if produced:
                return
        except Exception as exc:
            errors.append(f"OpenRouter: {exc}")

    if not CLOUD_MODE:
        yield from ollama.chat(
            model=local_model,
            messages=messages,
            stream=True,
            keep_alive="30m",
            options=model_options,
        )
        return

    if errors:
        raise RuntimeError("; ".join(errors)[-900:])
    raise RuntimeError("No cloud AI provider is configured.")


# =========================================================
# SIDEBAR
# =========================================================
with st.sidebar:
    st.markdown("## ✦ Nova")
    st.caption("Your AI workspace")

    if st.button("＋ New chat", use_container_width=True):
        new_id = create_chat()
        st.session_state.chat_id = new_id
        st.session_state.messages = []
        st.session_state.file_context = ""
        st.session_state.uploaded_files = []
        st.session_state.file_signature = ""
        st.session_state.generated_pdf = None
        st.session_state.generated_pdf_chat_id = None
        st.rerun()

    st.divider()
    st.markdown("### 💬 Chats")

    chats = get_chats()

    for cid, title in chats:
        col1, col2 = st.columns([5, 1])

        with col1:
            label = title if title else "New chat"
            if st.button(
                label[:38],
                key=f"chat_{cid}",
                use_container_width=True
            ):
                st.session_state.chat_id = cid
                st.session_state.messages = get_messages(cid)
                st.session_state.file_context = ""
                st.session_state.uploaded_files = []
                st.session_state.file_signature = ""
                st.session_state.generated_pdf = None
                st.session_state.generated_pdf_chat_id = None
                st.rerun()

        with col2:
            # Use a plain button because font-based dropdown icons can render
            # as raw text when an external icon font is blocked.
            if st.button("⋮", key=f"menu_{cid}", use_container_width=True):
                current_menu = st.session_state.get("open_chat_menu")
                st.session_state.open_chat_menu = (
                    None if current_menu == cid else cid
                )
                st.rerun()

        if st.session_state.get("open_chat_menu") == cid:
            with st.container(border=True):
                st.caption("Chat actions")
                if st.button(
                    "Delete chat",
                    key=f"confirm_delete_{cid}",
                    use_container_width=True,
                ):
                    delete_chat(cid)
                    st.session_state.open_chat_menu = None

                    if cid == st.session_state.chat_id:
                        new_id = create_chat()
                        st.session_state.chat_id = new_id
                        st.session_state.messages = []

                    st.rerun()

    use_internet = st.toggle("🌐 Internet", value=True)


# =========================================================
# MAIN CHAT
# =========================================================
if not st.session_state.messages:
    st.markdown("""
    <style>
    .chat-title {
        margin-top:105px !important;
        font-size:36px !important;
        font-weight:400 !important;
    }
    .chat-subtitle {display:none !important;}
    </style>
    """, unsafe_allow_html=True)
    st.markdown(
        '<div class="chat-title">What should we work on?</div>',
        unsafe_allow_html=True,
    )

    if st.session_state.uploaded_files:
        st.info(
            "📂 Ready: "
            + ", ".join(st.session_state.uploaded_files)
        )

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        render_chat_content(msg["content"])

# ChatGPT-like jump-to-latest control. The small component installs the button
# in the parent Streamlit page so it can stay fixed above the composer and keep
# working while streamed Markdown grows.
components.html(
    """
    <script>
    (() => {
      const w = window.parent;
      const d = w.document;
      const previous = w.__novaJumpToLatest;
      if (previous) {
        previous.cleanup();
      }

      const button = d.createElement('button');
      button.id = 'nova-jump-latest';
      button.type = 'button';
      button.setAttribute('aria-label', 'Jump to latest message');
      button.title = 'Jump to latest message';
      button.textContent = '↓';
      Object.assign(button.style, {
        position: 'fixed',
        left: 'calc(50% + 115px)',
        bottom: '108px',
        transform: 'translateX(-50%)',
        width: '38px',
        height: '38px',
        borderRadius: '999px',
        border: '1px solid #d8d8dc',
        background: '#ffffff',
        color: '#333333',
        fontSize: '22px',
        lineHeight: '34px',
        cursor: 'pointer',
        zIndex: '999999',
        boxShadow: '0 2px 10px rgba(0,0,0,.12)',
        display: 'none'
      });

      const mobile = w.matchMedia('(max-width: 768px)');
      const placeButton = () => {
        button.style.left = mobile.matches ? '50%' : 'calc(50% + 115px)';
        button.style.bottom = mobile.matches ? '142px' : '108px';
      };
      placeButton();

      const scrollTargets = () => Array.from(new Set([
        d.querySelector('[data-testid="stMain"]'),
        d.querySelector('[data-testid="stAppViewContainer"]'),
        d.scrollingElement,
        d.documentElement,
        d.body
      ].filter(Boolean)));

      const activeScroller = () => {
        const targets = scrollTargets();
        const movedTargets = targets.filter((target) => target.scrollTop > 0);
        const candidates = movedTargets.length ? movedTargets : targets;
        return candidates.sort((a, b) =>
          (b.scrollHeight - b.clientHeight) - (a.scrollHeight - a.clientHeight)
        )[0] || d.documentElement;
      };

      const distanceFromBottom = () => {
        const target = activeScroller();
        const pageDistance = Math.max(
          0,
          Math.max(d.body.scrollHeight, d.documentElement.scrollHeight)
            - (w.scrollY + w.innerHeight)
        );
        const targetDistance = Math.max(
          0,
          target.scrollHeight - target.scrollTop - target.clientHeight
        );
        return Math.max(pageDistance, targetDistance);
      };

      const hideStreamlitOwnerChrome = () => {
        d.querySelectorAll(
          '[data-testid="stAppDeployButton"], [class*="viewerBadge"], '
          + '[class*="ViewerBadge"]'
        ).forEach((node) => node.style.setProperty('display', 'none', 'important'));

        // Community Cloud injects this owner-only control outside Nova's
        // normal Streamlit component tree, so ordinary CSS selectors vary by
        // release. Match its exact visible label and hide its clickable shell.
        d.querySelectorAll('button, a, [role="button"], div').forEach((node) => {
          if ((node.textContent || '').trim() !== 'Manage app') return;
          const shell = node.closest('button, a, [role="button"]') || node;
          shell.style.setProperty('display', 'none', 'important');
          let parent = shell.parentElement;
          for (let depth = 0; parent && depth < 5; depth += 1) {
            if (w.getComputedStyle(parent).position === 'fixed') {
              parent.style.setProperty('display', 'none', 'important');
              break;
            }
            parent = parent.parentElement;
          }
        });
      };

      const refresh = () => {
        hideStreamlitOwnerChrome();
        d.querySelectorAll('[data-testid="stChatMessageContent"] a').forEach((link) => {
          link.setAttribute('target', '_blank');
          link.setAttribute('rel', 'noopener noreferrer');
        });
        const hasConversation = d.querySelectorAll('[data-testid="stChatMessage"]').length > 0;
        button.style.display = hasConversation && distanceFromBottom() > 180
          ? 'block' : 'none';
      };

      const jump = () => {
        const composer = d.querySelector('[data-testid="stChatInput"]');
        if (composer) {
          composer.scrollIntoView({behavior: 'smooth', block: 'end'});
        }
        scrollTargets().forEach((target) => {
          if (typeof target.scrollTo === 'function') {
            target.scrollTo({top: target.scrollHeight, behavior: 'smooth'});
          }
        });
        w.scrollTo({top: Math.max(d.body.scrollHeight, d.documentElement.scrollHeight), behavior: 'smooth'});
        w.setTimeout(refresh, 350);
      };

      button.addEventListener('click', jump);
      const listenedTargets = scrollTargets();
      listenedTargets.forEach((target) =>
        target.addEventListener('scroll', refresh, {passive: true})
      );
      w.addEventListener('scroll', refresh, {passive: true});
      mobile.addEventListener('change', placeButton);
      const observer = new MutationObserver(refresh);
      observer.observe(d.body, {childList: true, subtree: true, characterData: true});
      d.body.appendChild(button);
      refresh();

      w.__novaJumpToLatest = {
        cleanup: () => {
          observer.disconnect();
          listenedTargets.forEach((target) =>
            target.removeEventListener('scroll', refresh)
          );
          w.removeEventListener('scroll', refresh);
          mobile.removeEventListener('change', placeButton);
          button.remove();
        }
      };
    })();
    </script>
    """,
    height=0,
    width=0,
)

if (
    st.session_state.generated_pdf
    and st.session_state.generated_pdf_chat_id == st.session_state.chat_id
):
    st.download_button(
        "⬇️ Download PDF",
        data=st.session_state.generated_pdf,
        file_name=st.session_state.generated_pdf_name,
        mime="application/pdf",
        use_container_width=False,
        key=f"pdf_download_{st.session_state.chat_id}",
    )


# ChatGPT-like bottom controls: attachment plus model selector.
if GEMINI_API_KEY:
    MODEL_LABELS = {
        "gemini-3.6-flash": "Gemini 3.6 Flash",
        "gemini-3.1-flash-lite": "Gemini 3.1 Flash Lite",
        "gemini-2.5-flash-lite": "Gemini 2.5 Flash Lite",
        "gemini-2.5-flash": "Gemini 2.5 Flash",
    }
else:
    MODEL_LABELS = {
        "llama3.2": "Llama 3.2  Fast",
        "deepseek-r1:latest": "DeepSeek R1  Reasoning",
        "qwen2.5:latest": "Qwen 2.5  Balanced",
    }

def render_composer(container_key):
    """Render one native composer without any fixed-position CSS hacks."""
    with st.container(key=container_key):
        input_col, model_col = st.columns(
            [7.2, 2.2], vertical_alignment="bottom"
        )
        with input_col:
            submission = st.chat_input(
                "Message Nova...",
                accept_file="multiple",
                file_type=None,
                submit_mode="disable",
                key="nova_chat_input",
            )
        with model_col:
            selected_model = st.selectbox(
                "Model",
                list(MODEL_LABELS),
                format_func=lambda value: MODEL_LABELS[value],
                label_visibility="collapsed",
                key="selected_model",
            )
    return submission, selected_model


# One permanent composer for every state. New chat and existing chats use the
# exact same bottom position, matching ChatGPT behaviour.
with st.bottom:
    submission, model = render_composer("bottom_composer")


def warm_ollama_model(model_name):
    """Load the selected model in the background without blocking the UI."""
    try:
        ollama.chat(
            model=model_name,
            messages=[{"role": "user", "content": "hi"}],
            stream=False,
            keep_alive="30m",
            options={"num_predict": 1, "num_ctx": 2048},
        )
    except Exception:
        pass


if not CLOUD_MODE and st.session_state.get("warming_model") != model:
    st.session_state.warming_model = model
    threading.Thread(
        target=warm_ollama_model,
        args=(model,),
        daemon=True,
    ).start()


prompt = None
if submission:
    prompt = (submission.text or "").strip()
    submitted_files = list(submission.files or [])

    if submitted_files:
        sig = make_signature(submitted_files)
        if sig != st.session_state.file_signature:
            with st.spinner("Reading files..."):
                st.session_state.file_context = build_file_context(
                    submitted_files
                )
            st.session_state.file_signature = sig
            st.session_state.uploaded_files = [
                file.name for file in submitted_files
            ]

        if not prompt:
            prompt = "Analyze the uploaded file(s)."


if prompt:
    pdf_mode = get_pdf_request_mode(
        prompt,
        has_uploaded_file=bool(st.session_state.file_context),
    )
    pdf_requested = pdf_mode is not None
    resume_task = is_resume_request(prompt)
    coding_task = is_code_build_request(prompt)
    # Keep the attachment visible in chat history so users can verify that
    # Nova actually received the file used for the answer/document.
    attachment_names = [file.name for file in submitted_files]
    displayed_prompt = prompt
    if attachment_names:
        displayed_prompt += "\n\n📎 " + ", ".join(attachment_names)

    # Save user message immediately.
    save_message(st.session_state.chat_id, "user", displayed_prompt)
    st.session_state.messages.append(
        {"role": "user", "content": displayed_prompt}
    )

    # First user message becomes the sidebar chat title.
    user_messages = [
        m["content"]
        for m in st.session_state.messages
        if m["role"] == "user"
    ]

    if len(user_messages) == 1:
        title = prompt.strip().replace("\n", " ")
        update_chat_title(
            st.session_state.chat_id,
            title[:55] if title else "New chat"
        )

    with st.chat_message("user"):
        st.markdown(prompt)
        if attachment_names:
            safe_names = html.escape(", ".join(attachment_names))
            st.markdown(
                f'<div class="uploaded-file-note">📎 {safe_names}</div>',
                unsafe_allow_html=True,
            )

    # Image requests are tool actions, not questions for the text model. Nova
    # generates new artwork for creation prompts and searches the web only for
    # requests that ask for existing pictures.
    if wants_images(prompt):
        with st.chat_message("assistant"):
            if not use_internet:
                answer = "Please turn on the Internet toggle to use images."
                st.info(answer)
            elif wants_image_generation(prompt) and POLLINATIONS_API_KEY:
                try:
                    with st.spinner("Generating your image..."):
                        image_prompt, image_path = generate_ai_image(prompt)
                    answer = pack_generated_image(image_prompt, image_path)
                    render_chat_content(answer)
                except Exception:
                    # Never expose provider URLs/authentication failures. Give
                    # the user a useful clean gallery when generation is busy.
                    with st.spinner("Finding the closest matching pictures..."):
                        found_images = image_search(prompt)
                    if found_images:
                        answer = pack_image_results(
                            prompt, found_images, clean_gallery=True
                        )
                        render_chat_content(answer)
                    else:
                        answer = "I couldn't create or find a usable image right now."
                        st.warning(answer)
            else:
                try:
                    with st.spinner("Finding pictures..."):
                        found_images = image_search(prompt)
                    if found_images:
                        answer = pack_image_results(
                            prompt, found_images, clean_gallery=True
                        )
                        render_chat_content(answer)
                    else:
                        answer = "I couldn't find usable image results for that search."
                        st.warning(answer)
                except Exception:
                    answer = "Image search is temporarily unavailable. Please retry shortly."
                    st.error(answer)

        save_message(st.session_state.chat_id, "assistant", answer)
        st.session_state.messages.append(
            {"role": "assistant", "content": answer}
        )
        st.rerun()

    # Current weather is a deterministic tool call. It does not consume the
    # shared Gemini quota, so it remains reliable when several people use Nova.
    if wants_weather(prompt):
        with st.chat_message("assistant"):
            if not use_internet:
                answer = (
                    "Please turn on the Internet toggle to get live weather."
                )
                st.info(answer)
            else:
                try:
                    with st.spinner("Checking live weather..."):
                        answer = get_live_weather(prompt)
                    st.markdown(answer)
                except Exception as exc:
                    answer = (
                        "I couldn't fetch live weather right now. Please try "
                        f"again shortly. ({html.escape(str(exc))[:180]})"
                    )
                    st.warning(answer)

        save_message(st.session_state.chat_id, "assistant", answer)
        st.session_state.messages.append(
            {"role": "assistant", "content": answer}
        )
        st.rerun()

    # Common greetings should respond instantly without waking the full LLM.
    simple_prompt = re.sub(r"[^a-z ]", "", prompt.lower()).strip()
    instant_replies = {
        "hi": "Hi! How can I help you?",
        "hello": "Hello! How can I help you?",
        "hey": "Hey! How can I help you?",
        "hi nova": "Hi! How can I help you?",
        "hello nova": "Hello! How can I help you?",
        "hey nova": "Hey! How can I help you?",
    }
    if simple_prompt in instant_replies:
        answer = instant_replies[simple_prompt]
        with st.chat_message("assistant"):
            st.markdown(answer)
        save_message(st.session_state.chat_id, "assistant", answer)
        st.session_state.messages.append(
            {"role": "assistant", "content": answer}
        )
        st.rerun()

    date_prompts = {
        "what is today date", "what is todays date", "today date",
        "aaj ki date", "aaj date kya hai",
    }
    if simple_prompt in date_prompts:
        answer = datetime.now().strftime("Today is %d %B %Y.")
        with st.chat_message("assistant"):
            st.markdown(answer)
        save_message(st.session_state.chat_id, "assistant", answer)
        st.session_state.messages.append(
            {"role": "assistant", "content": answer}
        )
        st.rerun()

    # PDF requests are handled directly by the app. Do not ask the LLM to
    # rewrite the document inside the chat before creating the file.
    if pdf_mode == "export":
        with st.chat_message("assistant"):
            try:
                source_messages = st.session_state.messages[:-1]

                if not source_messages:
                    raise ValueError(
                        "There is no previous conversation to export yet."
                    )

                pdf_bytes = build_chat_pdf(
                    source_messages,
                    title="Nova - Questions and Answers",
                )
                st.session_state.generated_pdf = pdf_bytes
                st.session_state.generated_pdf_name = (
                    "Nova_Questions_and_Answers.pdf"
                )
                st.session_state.generated_pdf_chat_id = (
                    st.session_state.chat_id
                )

                ready_message = "✅ Your PDF is ready."
                st.success(ready_message)
                st.download_button(
                    "⬇️ Download PDF",
                    data=pdf_bytes,
                    file_name=st.session_state.generated_pdf_name,
                    mime="application/pdf",
                    key=f"new_pdf_{st.session_state.chat_id}",
                )
                save_message(
                    st.session_state.chat_id,
                    "assistant",
                    ready_message,
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": ready_message}
                )
            except Exception as pdf_error:
                error_message = f"❌ Could not create PDF: {pdf_error}"
                st.error(error_message)
                save_message(
                    st.session_state.chat_id,
                    "assistant",
                    error_message,
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": error_message}
                )

        st.rerun()

    system_prompt = """
You are Nova, a fast and helpful AI assistant.

The application's current date is {current_date}.

Rules:
- Be accurate, practical and concise.
- Format answers like a polished ChatGPT response using valid Markdown.
- Use short paragraphs, descriptive headings only when useful, clean bullet or
  numbered lists, fenced code blocks with a language tag, and Markdown tables
  for genuine comparisons. Do not over-format a simple one-line answer.
- Never leave an incomplete bullet, sentence, heading or section.
- Remember the conversation context provided below.
- Treat the newest user message as the controlling request. Never silently
  carry a company, location, date, role or other filter from an older message
  when the newest message asks generally, broadly, for all, or changes topic.
- In particular, an earlier discussion about Amex does not mean every later
  Data Analyst or career question is about Amex. Retain Amex only when the
  newest message explicitly says Amex or clearly refers back to that company.
- If a file is uploaded, use the file as the primary source.
- Never invent values that are not present in the uploaded file.
- For Excel/CSV questions, reason from the supplied data.
- For SQL/Python questions, provide correct practical code.
- If web results are supplied, use them for current information.
- WEB SEARCH RESULTS are fresh external evidence and override older model
  knowledge whenever they conflict.
- When WEB SEARCH RESULTS are present, answer the question directly from them.
  Honor the exact count requested in the current message or follow-up (for
  example, return 10 complete items when the user corrects 4 to 10).
  Never mention a knowledge-cutoff date and never say that your knowledge only
  goes up to 2023.
- If fresh results are insufficient, say only what could not be verified; do
  not present the model cutoff as a weakness.
- Do not claim to have browsed unless WEB SEARCH RESULTS are actually supplied.
- Every external Markdown link must use a URL copied exactly from supplied WEB
  SEARCH RESULTS. Never guess, shorten, reconstruct or invent a careers URL.
- If the user reports that a link does not open, use fresh search results and
  provide the best verified direct link instead of speculating about cookies,
  ad blockers or regional redirects.
- If DOWNLOADED FILE CONTENT is supplied, the application has already downloaded and read the file. Analyze it directly.
- Do not claim that you cannot download files when DOWNLOAD STATUS is supplied.
- Never invent values that are not present in supplied file data.
- Never write fake links, placeholder links, "insert downloadable link", or
  claim that a PDF/file exists. The application itself creates download files.
""".format(current_date=datetime.now().strftime("%d %B %Y"))

    if coding_task:
        system_prompt += """

COMPLETE CODING / APP-BUILD TASK:
- Build the requested working project now. Do not refuse, stall, give only an
  outline, or ask unnecessary follow-up questions when reasonable defaults can
  be chosen.
- Return a minimal but complete runnable implementation, not pseudocode or
  isolated snippets. Include every required file in a separate fenced code
  block whose heading is the exact filename.
- Include setup/run commands and a short project tree before the files.
- Preserve the user's requested language/framework. If none is specified, use
  the simplest browser-runnable HTML/CSS/JavaScript solution; use placeholder
  media/assets where copyrighted or private service assets are unavailable.
- For a basic clone, reproduce the core layout and interactions without
  copying trademarks, credentials, private APIs or copyrighted media.
- Check imports, variable names, event handlers and file references silently
  before answering. Never leave TODOs, ellipses, incomplete code, or say
  "continue" for the remaining files.
"""

    if st.session_state.file_context and (
        pdf_mode == "content" or should_use_file_context(prompt)
    ):
        file_context_for_prompt = st.session_state.file_context
        if pdf_mode == "content" and len(file_context_for_prompt) > 18000:
            file_context_for_prompt = (
                file_context_for_prompt[:18000]
                + "\n[Source shortened for faster PDF generation.]"
            )
        system_prompt += (
            "\n\nUPLOADED FILE CONTENT:\n"
            + file_context_for_prompt
        )

    if resume_task:
        system_prompt += """

CV / RESUME TASK (APPLIES IN CHAT AND PDF MODE):
- Rewrite the complete uploaded CV according to the user's requested changes.
- Preserve the candidate's real name, contact details, employers, dates,
  education, skills, projects and achievements from the source.
- Never invent experience, employers, qualifications, metrics or tools.
- Never change the candidate's employer, job title or employment history to
  match the target role. Reposition only verified skills and projects.
- If the source employer is Bharti Airtel, it must remain Bharti Airtel.
- Never introduce Celebal Technologies or any other employer absent from the
  uploaded source.
- Never add RAG, LangChain, Vector Databases, Azure OpenAI, CrewAI, FastAPI,
  Docker, MLflow, PromptFlow, LangSmith, XGBoost, Prophet, accuracy figures,
  cost reductions or document counts unless explicitly present in the source.
- Do not use unsupported promotional claims such as "expert", "sophisticated",
  "deep analysis", "enterprise-grade", "real-time", "high-volume",
  "architected" or "significantly improved" unless the uploaded CV explicitly
  proves the claim.
- Improve wording, grammar, impact and ATS compatibility.
- Return a complete ready-to-use resume, not advice, examples, formulas,
  explanations or placeholders.
- Use clear resume sections and concise achievement-oriented bullet points.
- Use this exact Markdown contract so both chat and the PDF are professionally
  formatted: one `# CANDIDATE NAME` line; one plain contact line; `##` section
  headings; `###` role/project/education entries; and `-` for every bullet.
- Use sections in this order unless the user requests otherwise: PROFESSIONAL
  SUMMARY, CORE SKILLS, PROFESSIONAL EXPERIENCE, PROJECTS, EDUCATION.
- Keep headings, employer names, job titles and dates consistent. Do not output
  horizontal rules, tables, block quotes, emojis, raw HTML or decorative text.
- When the target is Data Analyst / GenAI, include this verified project as a
  flagship project (these are facts about this running application):
  "Nova - Deployed GenAI Data Analyst Assistant" built with Python,
  Streamlit, Gemini API, Ollama, SQLite, Pandas and document parsers. It
  supports PDF/Excel/CSV/DOCX/JSON/text analysis, context-window memory,
  session-isolated chat history, intent routing, web/image search, streaming
  responses, secure Streamlit Secrets and PDF generation.
- Include the verified links:
  Live Demo: https://nova-ai-chatbot-4sev2pjpwh9nimi23pjmhn.streamlit.app
  GitHub: https://github.com/Tanya884/nova-ai-chatbot
- If the uploaded candidate is Tanya Tiwari, ensure the final resume includes
  these verified facts in addition to the source CV:
  * 3 years of experience based on the Jul 2023 - Present employment dates.
  * Partner Performance reporting across Nodes, Backbone, KM, ATC, NC, SD/BG,
    Capping and DLP; preparation time reduced from about 20 minutes to 1.5
    minutes per partner.
  * Power BI weekly operational reporting for ring performance, upgrades,
    high-utilization sites and material inventory.
  * Patroller Location Validation using an image-analysis workflow for 1,000+
    field images to check whether a patroller is within a defined range.
  * KMZ Node Location Extraction to derive latitude/longitude coordinates.
- For Tanya Tiwari, order the resume as: Professional Summary, Core Skills,
  Professional Experience, Nova flagship project, other selected projects,
  Education. Keep it to two pages and use concise ATS-friendly bullets.
- Before responding, silently verify that every required Tanya project and the
  two Nova links are present and that no unsupported claim was added.
- Output the complete final resume in one response. Never stop midway, ask the
  user to type "yes", or provide only a partial section.
- Output only the final resume content, without introductory advice.
"""
    elif pdf_mode == "content":
        system_prompt += """

PDF DOCUMENT TASK:
- Create the complete document requested by the user using the uploaded file
  and conversation as the source.
- Follow the user's requested document type and requested changes exactly.
- Cover the important source content and organize it with clear headings.
- If the user specifically asks for interview Q&A, include complete questions
  with concise interview-ready answers.
- Output only the document content. Do not mention links, downloads, PDF-style
  text, limitations, or ask follow-up questions.
The application will privately convert your output into a real PDF.
"""

    # Short follow-ups such as "10 chahiye", "aur do" and "batao" inherit
    # the most recent live-search request instead of becoming isolated prompts.
    recent_web_seed = ""
    for old_message in reversed(st.session_state.messages[:-1][-10:]):
        old_text = str(old_message.get("content", ""))
        if old_message.get("role") == "user" and needs_web_search(old_text):
            recent_web_seed = old_text
            break

    followup_terms = [
        "batao", "btao", "aur", "more", "chahiye", "chaye", "only",
        "but", "yeh", "this", "same", "continue", "fir", "phir",
    ]
    short_followup = len(prompt.split()) <= 14 and any(
        term in prompt.lower() for term in followup_terms
    )
    inherited_web = bool(short_followup and recent_web_seed)
    web_needed = use_internet and (
        needs_web_search(prompt) or inherited_web
    )
    web_query = (
        recent_web_seed + "\nFollow-up requirement: " + prompt
        if inherited_web
        else prompt
    )

    fact_check = needs_fact_verification(web_query)
    if fact_check:
        web_query += (
            "\nFind the exact figure from authoritative or official sources. "
            "Cross-check the number and definition before answering."
        )

    results = ""
    if web_needed:
        results = web_search(web_query)

        if fact_check:
            system_prompt += """

FACT-CHECKING REQUIREMENT:
- The user is asking for an exact public fact or statistic. Use only the
  supplied WEB SEARCH RESULTS; do not answer from memory.
- Prefer government, census, regulator, university or other primary sources.
- Check whether the requested category, geography, year and unit match the
  source. Do not combine a count from one definition with a percentage from
  another.
- State the source name and year next to the answer. If the results do not
  verify one exact figure, clearly say that instead of estimating.
"""

        if wants_download(prompt):
            downloaded_path, download_status = try_download_from_results(results)

            if downloaded_path:
                try:
                    import pandas as pd

                    ext = os.path.splitext(downloaded_path.lower())[1]

                    if ext == ".csv":
                        downloaded_content = pd.read_csv(
                            downloaded_path
                        ).to_string(index=False)

                    elif ext in [".xlsx", ".xls", ".xlsm"]:
                        book = pd.ExcelFile(downloaded_path)
                        parts = []
                        for sheet in book.sheet_names:
                            df = pd.read_excel(downloaded_path, sheet_name=sheet)
                            parts.append(f"--- Excel Sheet: {sheet} ---")
                            parts.append(df.to_string(index=False))
                        downloaded_content = "\n".join(parts)

                    elif ext == ".json":
                        downloaded_content = json.dumps(
                            json.load(open(downloaded_path, encoding="utf-8")),
                            indent=2,
                            ensure_ascii=False
                        )

                    else:
                        downloaded_content = Path(downloaded_path).read_text(
                            encoding="utf-8", errors="replace"
                        )

                    if len(downloaded_content) > 60000:
                        downloaded_content = (
                            downloaded_content[:60000]
                            + "\n[Downloaded file content truncated.]"
                        )

                    system_prompt += (
                        "\n\nDOWNLOAD STATUS:\n"
                        + download_status
                        + "\n\nDOWNLOADED FILE CONTENT:\n"
                        + downloaded_content
                    )

                except Exception as e:
                    system_prompt += (
                        "\n\nDOWNLOAD STATUS:\n"
                        + download_status
                        + f"\nCould not parse downloaded file: {e}"
                    )
            else:
                system_prompt += "\n\nDOWNLOAD STATUS:\n" + download_status

        system_prompt += "\n\nWEB SEARCH RESULTS:\n" + results

    # PDF generation needs the uploaded source and current request, not the
    # entire chat history. This greatly reduces prompt-processing time.
    if pdf_mode == "content":
        window_messages = st.session_state.messages[-1:]
    elif requests_broad_scope(prompt):
        # "All positions/companies" deliberately removes an earlier company
        # filter (for example Amex), so isolate the current request.
        window_messages = st.session_state.messages[-1:]
    else:
        window_messages = st.session_state.messages[-WINDOW_MEMORY:]

    nova_messages = [
        {"role": "system", "content": system_prompt}
    ] + window_messages

    with st.chat_message("assistant"):
        placeholder = st.empty()
        answer = ""
        placeholder.markdown("Thinking…")

        try:
            model_options = {
                "temperature": 0.2 if resume_task else TEMPERATURE,
                # Give both normal chat and generated PDFs the same generous
                # response budget. The provider/model may enforce a lower
                # hard maximum, but Nova must not impose the old 1.1k/5k cap.
                "num_ctx": 32768,
                "num_predict": 25000,
            }
            if resume_task:
                model_options.update({
                    "num_batch": 512,
                })
            elif pdf_mode == "content":
                model_options.update({
                    "num_batch": 512,
                })

            stream = stream_nova_chat(
                local_model=model,
                messages=nova_messages,
                model_options=model_options,
            )

            writing_started = False
            for chunk in stream:
                piece = chunk.get("message", {}).get(
                    "content", ""
                )

                answer += piece
                if pdf_mode == "content" and piece and not writing_started:
                    placeholder.markdown("Writing your PDF…")
                    writing_started = True
                elif pdf_mode != "content":
                    placeholder.markdown(answer + "▌")

            # Provider-level retries and alternate models are already handled
            # inside stream_nova_chat. A second full call here would only burn
            # more of the shared free-tier quota.
            if not answer.strip():
                answer = build_search_fallback(web_query, results)
                if not answer:
                    raise RuntimeError(
                        "All configured AI providers are temporarily busy."
                    )

            if pdf_mode == "content":
                if resume_task:
                    pdf_bytes = build_resume_pdf(answer)
                else:
                    pdf_bytes = build_chat_pdf(
                        [{"role": "assistant", "content": answer}],
                        title="Nova - Generated Document",
                        show_role_labels=False,
                    )
                st.session_state.generated_pdf = pdf_bytes
                st.session_state.generated_pdf_name = (
                    "Tanya_Tiwari_Resume.pdf"
                    if resume_task
                    else "Nova_Generated_Document.pdf"
                )
                st.session_state.generated_pdf_chat_id = (
                    st.session_state.chat_id
                )

                ready_message = "✅ Your PDF is ready."
                placeholder.success(ready_message)
                st.download_button(
                    "⬇️ Download PDF",
                    data=pdf_bytes,
                    file_name=st.session_state.generated_pdf_name,
                    mime="application/pdf",
                    key=f"content_pdf_{st.session_state.chat_id}",
                )
                save_message(
                    st.session_state.chat_id,
                    "assistant",
                    ready_message,
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": ready_message}
                )
            else:
                placeholder.markdown(answer)
                save_message(
                    st.session_state.chat_id,
                    "assistant",
                    answer,
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": answer}
                )

        except Exception as e:
            fallback_answer = (
                build_search_fallback(web_query, results)
                if web_needed and pdf_mode != "content"
                else ""
            )
            if fallback_answer:
                placeholder.markdown(fallback_answer)
                save_message(
                    st.session_state.chat_id,
                    "assistant",
                    fallback_answer,
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": fallback_answer}
                )
            elif GEMINI_API_KEY:
                error_message = (
                    "❌ Nova is temporarily busy because the shared free "
                    "AI quota is full. Please retry in a minute."
                )
                placeholder.error(error_message)
                save_message(
                    st.session_state.chat_id, "assistant", error_message
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": error_message}
                )
            elif OPENROUTER_API_KEY:
                error_message = (
                    "❌ Nova could not reach the cloud model. Check the "
                    "OpenRouter key, balance, and model access in your "
                    "deployment settings."
                )
                placeholder.error(error_message)
                save_message(
                    st.session_state.chat_id, "assistant", error_message
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": error_message}
                )
            else:
                error_message = f"❌ Error: {e}"
                placeholder.error(error_message)
                save_message(
                    st.session_state.chat_id, "assistant", error_message
                )
                st.session_state.messages.append(
                    {"role": "assistant", "content": error_message}
                )

    # Re-render once after completion. This removes the temporary new-chat
    # layout and leaves the finished conversation in the compact chat view.
    st.rerun()
